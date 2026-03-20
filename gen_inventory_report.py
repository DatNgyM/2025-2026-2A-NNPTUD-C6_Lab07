import os
import textwrap
import datetime

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


BASE = r"E:\2025-2026-2A-NNPTUD-C6_Lab07"
DOCS_DIR = os.path.join(BASE, "docs")
IMG_DIR = os.path.join(DOCS_DIR, "postman_images")


def get_font(size: int):
    # Try common fonts available on Windows/Linux.
    for name in ["arial.ttf", "Arial.ttf", "DejaVuSans.ttf", "DejaVuSansMono.ttf"]:
        try:
            return ImageFont.truetype(name, size)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap_lines(text: str, width: int = 80):
    lines = []
    if not text.strip():
        return lines
    for paragraph in text.split("\n"):
        if paragraph.strip() == "":
            lines.append("")
            continue
        wrapped = textwrap.wrap(paragraph, width=width)
        lines.extend(wrapped if wrapped else [""])
    return lines


def draw_text_multiline(draw: ImageDraw.ImageDraw, text: str, x: int, y: int, font, fill, max_y: int):
    line_h = getattr(font, "size", 22) + 8
    for line in wrap_lines(text, width=80):
        if y > max_y:
            break
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h
    return y


def generate_postman_image(ep: dict, out_path: str):
    W, H = 1280, 720
    img = Image.new("RGB", (W, H), (30, 30, 30))
    draw = ImageDraw.Draw(img)

    FONT_HDR = get_font(24)
    FONT_TITLE = get_font(34)
    FONT_BODY = get_font(22)
    FONT_SMALL = get_font(18)

    # Header
    draw.rectangle([0, 0, W, 90], fill=(45, 45, 45))
    draw.text((30, 20), ep["method"], font=FONT_HDR, fill=(255, 255, 255))
    draw.text((160, 20), ep["url"], font=FONT_HDR, fill=(180, 220, 255))

    # Left panel (request body)
    draw.rectangle([0, 90, 520, H], fill=(32, 32, 32))
    draw.text((30, 110), "Body", font=FONT_SMALL, fill=(200, 200, 200))
    if ep["body"].strip():
        draw_text_multiline(draw, ep["body"], 30, 150, FONT_BODY, (220, 255, 220), H - 30)
    else:
        draw.text((30, 150), "(none)", font=FONT_SMALL, fill=(200, 200, 200))

    # Right panel (response)
    draw.rectangle([520, 90, W, H], fill=(25, 25, 25))
    draw.text((560, 110), "JSON Response", font=FONT_SMALL, fill=(200, 200, 200))
    draw_text_multiline(draw, ep["response"], 560, 150, FONT_BODY, (230, 230, 230), H - 30)

    img.save(out_path)


def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    os.makedirs(IMG_DIR, exist_ok=True)

    pid = "<productId>"
    iid = "<inventoryId>"

    endpoints = [
        {
            "name": "GET all inventories (join product)",
            "method": "GET",
            "url": "/api/v1/inventories",
            "body": "",
            "response": (
                '{ "_id": "' + iid + '", "product": { "_id": "' + pid + '", "title": "prod1" }, '
                '"stock": 10, "reserved": 2, "soldCount": 1 }'
            ),
        },
        {
            "name": "GET inventory by ID (join product)",
            "method": "GET",
            "url": "/api/v1/inventories/" + iid,
            "body": "",
            "response": (
                '{ "_id": "' + iid + '", "product": { "_id": "' + pid + '", "title": "prod1" }, '
                '"stock": 10, "reserved": 2, "soldCount": 1 }'
            ),
        },
        {
            "name": "POST add_stock",
            "method": "POST",
            "url": "/api/v1/inventories/add_stock",
            "body": '{ "product": "' + pid + '", "quantity": 10 }',
            "response": '{ "_id": "' + iid + '", "stock": 20, "reserved": 2, "soldCount": 1 }',
        },
        {
            "name": "POST remove_stock",
            "method": "POST",
            "url": "/api/v1/inventories/remove_stock",
            "body": '{ "product": "' + pid + '", "quantity": 5 }',
            "response": '{ "_id": "' + iid + '", "stock": 15, "reserved": 2, "soldCount": 1 }',
        },
        {
            "name": "POST reservation",
            "method": "POST",
            "url": "/api/v1/inventories/reservation",
            "body": '{ "product": "' + pid + '", "quantity": 3 }',
            "response": '{ "_id": "' + iid + '", "stock": 12, "reserved": 5, "soldCount": 1 }',
        },
        {
            "name": "POST sold",
            "method": "POST",
            "url": "/api/v1/inventories/sold",
            "body": '{ "product": "' + pid + '", "quantity": 2 }',
            "response": '{ "_id": "' + iid + '", "stock": 12, "reserved": 3, "soldCount": 3 }',
        },
    ]

    img_paths = []
    for i, ep in enumerate(endpoints, start=1):
        safe = ep["name"].split("(")[0].strip().replace(" ", "_").replace("/", "-")
        out = os.path.join(IMG_DIR, f"{i:02d}_{safe}.png")
        generate_postman_image(ep, out)
        img_paths.append(out)

    report_path = os.path.join(DOCS_DIR, "inventory_postman_report.docx")
    doc = Document()
    doc.add_heading("Inventory Management - Postman Functions Report", level=1)
    doc.add_paragraph("Ngày tạo: " + str(datetime.date.today()))
    doc.add_paragraph(
        "Lưu ý: Vì máy hiện chưa kết nối được MongoDB local để chụp response thật từ Postman, "
        "các hình dưới đây là minh hoạ payload/response theo đúng logic nghiệp vụ bạn yêu cầu. "
        "Bạn có thể thay bằng screenshot thực tế sau khi chạy API."
    )

    doc.add_heading("Các endpoint", level=2)
    for ep, img_path in zip(endpoints, img_paths):
        doc.add_heading(ep["name"], level=3)
        if ep["body"].strip():
            doc.add_paragraph("Request body: " + ep["body"])
        doc.add_paragraph("Response mẫu: " + ep["response"])
        doc.add_picture(img_path, width=Inches(6.5))
        doc.add_paragraph("")

    doc.add_heading("Tóm tắt URL", level=2)
    summary = [
        "GET  /api/v1/inventories",
        "GET  /api/v1/inventories/:id",
        "POST /api/v1/inventories/add_stock { product, quantity }",
        "POST /api/v1/inventories/remove_stock { product, quantity }",
        "POST /api/v1/inventories/reservation { product, quantity }",
        "POST /api/v1/inventories/sold { product, quantity }",
    ]
    for s in summary:
        doc.add_paragraph(s)

    doc.save(report_path)
    print("Created:", report_path)
    print("Created images:", len(img_paths))


if __name__ == "__main__":
    main()

